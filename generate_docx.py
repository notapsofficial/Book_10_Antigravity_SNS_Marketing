import os
import glob
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class ManuscriptBuilder:
    def __init__(self, filename="Antigravity_SNS_Marketing_Rev_01_JP.docx"):
        self.doc = Document()
        self.filename = filename
        self.output_dir = "manuscript"
        self._setup_styles()

    def _setup_styles(self):
        # Base style settings (trying to approximate Hiragino Mincho)
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Hiragino Mincho ProN'
        font.size = Pt(10.5)
        # For Japanese fonts in Word
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Mincho ProN')

    def add_title_page(self, title, author):
        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(24)
        
        # Spacer
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Author
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(author)
        run.font.size = Pt(14)

        # Page break
        self.doc.add_page_break()

    def parse_markdown_line(self, line):
        line = line.strip()
        if not line:
            return # Skip empty lines for now, or add small spacer

        # Headers
        header_match = re.match(r'^(#+)\s+(.*)', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            # In Word, Heading 1 is usually Title level, but here we map # -> Heading 1
            # Adjust level if needed. 
            # Docx supports Heading 1-9
            self.doc.add_heading(text, level=min(level, 9))
            return

        # List items
        list_match = re.match(r'^-\s+(.*)', line)
        if list_match:
            text = list_match.group(1)
            p = self.doc.add_paragraph(style='List Bullet')
            self._add_formatted_run(p, text)
            return
            
        # Images (Basic support) - ![alt](path)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            # image_path = img_match.group(2)
            # For now, just adding a placeholder text because we might not have the images or paths might be absolute/relative issues
            # In a real script we would load the image. 
            # self.doc.add_picture(image_path)
            self.doc.add_paragraph(f"[IMAGE: {line}]") # Placeholder
            return

        # Blockquotes >
        quote_match = re.match(r'^>\s+(.*)', line)
        if quote_match:
            text = quote_match.group(1)
            p = self.doc.add_paragraph(style='Quote') # Or Normal with indent
            self._add_formatted_run(p, text)
            return

        # Normal text
        p = self.doc.add_paragraph()
        self._add_formatted_run(p, line)

    def _add_formatted_run(self, paragraph, text):
        # Handle **Bold**
        # Split by **
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    def process_file(self, filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for line in lines:
            self.parse_markdown_line(line)
        
        # Add page break after each chapter
        self.doc.add_page_break()

    def save(self):
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
        path = os.path.join(self.output_dir, self.filename)
        self.doc.save(path)
        print(f"Saved to {path}")

def main():
    builder = ManuscriptBuilder()
    builder.add_title_page("ANTIGRAVITY × 書かないSNS\nAPIで実装する、個人開発者の発信パイプライン", "D.D.Scotick")

    # Define chapter order
    chapters = [
        "manuscript/chapter_00.md",
        "manuscript/chapter_01.md",
        "manuscript/chapter_02.md",
        "manuscript/chapter_03.md",
        "manuscript/chapter_04.md",
        "manuscript/chapter_05.md",
        "manuscript/chapter_06.md",
        "manuscript/chapter_07.md",
        "manuscript/chapter_08.md",
        "manuscript/chapter_09.md",
        "manuscript/chapter_10.md",
        "manuscript/chapter_11.md",
        "manuscript/chapter_12.md",
        "manuscript/chapter_13.md",
        "manuscript/chapter_14.md",
        "manuscript/chapter_15.md",
    ]

    for chapter_path in chapters:
        if os.path.exists(chapter_path):
            print(f"Processing {chapter_path}...")
            builder.process_file(chapter_path)
        else:
            print(f"Warning: {chapter_path} not found.")

    builder.save()

if __name__ == "__main__":
    main()
